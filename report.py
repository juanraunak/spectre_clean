import json
import os
from pathlib import Path
import time
from typing import Dict, Any, Optional
from openai import AzureOpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class SkillBoostPlanGenerator:
    def __init__(self):
        """Initialize with Azure OpenAI configuration (hard-coded)"""
        # Hardcoded Azure OpenAI configuration
        self.api_key = "2be1544b3dc14327b60a870fe8b94f35"
        self.endpoint = "https://notedai.openai.azure.com"
        self.api_version = "2024-06-01"
        self.deployment_id = "gpt-4o"
        
        # Initialize Azure OpenAI client
        self.client = AzureOpenAI(
            api_key=self.api_key,
            api_version=self.api_version,
            azure_endpoint=self.endpoint
        )

    def load_course_data(self, file_path: str) -> list:
        """Load course data from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return json.load(file)
        except Exception as e:
            print(f"Error loading file: {e}")
            return []
    
    def load_skill_gaps_data(self, file_path: str = "final_skill_gaps_detailed_gpt.json") -> dict:
        """Load skill gaps data from JSON file and create a lookup dictionary"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                skill_gaps_list = json.load(file)
                # Create lookup dictionary by employee name
                skill_gaps_dict = {}
                for employee in skill_gaps_list:
                    employee_name = employee.get('manipal_employee', '')
                    skill_gaps_dict[employee_name] = employee
                return skill_gaps_dict
        except Exception as e:
            print(f"Error loading skill gaps file: {e}")
            return {}
    
    def find_skill_gap_data(self, employee_name: str, skill_gaps_data: dict) -> Optional[Dict[str, Any]]:
        """Find skill gap data for a specific employee"""
        return skill_gaps_data.get(employee_name)
    
    def generate_skill_boost_plan(self, employee_data: Dict[str, Any], skill_gap_data: Optional[Dict[str, Any]] = None) -> str:
        """Generate enhanced skill boost plan using both course and skill gap data"""
        
        # Build enhanced prompt with skill gap data if available
        skill_gap_section = ""
        if skill_gap_data:
            competitor_info = f"Competitor companies analyzed: {', '.join(skill_gap_data.get('competitor_companies', []))}"
            competitor_count = skill_gap_data.get('competitor_count', 0)
            critical_gaps = []
            important_gaps = []
            
            for skill, importance in skill_gap_data.get('skill_importance', {}).items():
                if importance == "Critical":
                    critical_gaps.append(skill)
                elif importance in ["Important", "Nice-to-have"]:
                    important_gaps.append(skill)
            
            skill_gap_section = f"""
        
        SKILL GAP ANALYSIS (Based on Competitor Benchmarking):
        Total Competitors Analyzed: {competitor_count}
        {competitor_info}
        
        Critical Skill Gaps: {', '.join(critical_gaps) if critical_gaps else 'None identified'}
        Additional Skill Gaps: {', '.join(important_gaps) if important_gaps else 'None identified'}
        
        Gap Reasoning:
        {chr(10).join([f"- {skill}: {reasoning}" for skill, reasoning in skill_gap_data.get('gap_reasoning', {}).items()])}
        
        Overall Assessment: {skill_gap_data.get('overall_assessment', 'No assessment available')}
        Recommendations: {chr(10).join(skill_gap_data.get('recommendations', []))}
        """

        prompt = f"""
        Convert the following employee course data into a comprehensive 0-3 Month Skill Boost Plan report format.

        Employee Data:
        Name: {employee_data['employeeName']}
        Role: {employee_data['role']}
        Company: {employee_data['company']}
        Existing Skills: {', '.join(employee_data['existingSkills'])}
        Priority Skills Selected: {', '.join(employee_data['prioritySkillsSelected'])}
        
        Course Information:
        Course Name: {employee_data['course']['courseName']}
        Description: {employee_data['course']['description']}
        Skills Covered: {employee_data['course']['skillsCovered']}
        Total Topics: {employee_data['course']['totalTopics']}
        
        {skill_gap_section}

        Create a comprehensive report following this EXACT format:

        # {employee_data['employeeName']} -- 0-3 Month Skill Boost Plan

        This document outlines the 0-3 month skill upgrade plan for {employee_data['employeeName']}, focusing on closing immediate skill gaps, recommended courses, proof-of-concept (POC) deliverables, and the direct business impact for {employee_data['company']}.

        ## Executive Summary
        
        **Employee**: {employee_data['employeeName']}  
        **Current Role**: {employee_data['role']}  
        **Company**: {employee_data['company']}  
        **Assessment Date**: {time.strftime('%B %Y')}
        
        ### Skill Gap Analysis Overview
        {"- **Competitors Analyzed**: " + str(skill_gap_data.get('competitor_count', 0)) + " companies including " + ', '.join(skill_gap_data.get('competitor_companies', [])) if skill_gap_data else "- **Skill Gap Analysis**: Based on internal assessment"}
        {"- **Critical Skill Gaps Identified**: " + str(len([s for s, i in skill_gap_data.get('skill_importance', {}).items() if i == 'Critical'])) if skill_gap_data else ""}
        {"- **Confidence Score**: " + str(skill_gap_data.get('evidence_flags', {}).get('confidence_score', 'N/A')) if skill_gap_data else ""}
        - **Priority Skills for Development**: {', '.join(employee_data['prioritySkillsSelected'])}
        - **Course Modules**: {employee_data['course']['totalTopics']} topics across {len(employee_data['course']['course'])} chapters

        ## Detailed Learning Roadmap

        | Timeline | Skill Gap | Priority Level | Recommended Skills/Tools | Course Recommendation | Proposed POC | Business Impact for {employee_data['company']} | Competitor Benchmark |
        |----------|-----------|----------------|-------------------------|----------------------|--------------|-------------------------------------------|-------------------|
        | 0-1 Month | [First critical skill gap] | Critical | [Specific tools/technologies] | [Course name and key modules] | [Practical deliverable] | [Direct business value] | [Competitor examples] |
        | 0-1 Month | [Second critical skill gap] | Critical | [Specific tools/technologies] | [Course name and key modules] | [Practical deliverable] | [Direct business value] | [Competitor examples] |
        | 1-3 Months | [Important skill gap] | Important | [Advanced tools/technologies] | [Advanced course modules] | [Complex deliverable] | [Strategic business value] | [Competitor examples] |
        | 1-3 Months | [Another important skill] | Nice-to-have | [Specialized tools] | [Specialized modules] | [Industry-specific POC] | [Long-term business impact] | [Competitor examples] |

        ## Execution Priority

        ### Phase 1: Weeks 1-4 (Critical Skills)
        Focus on [priority skills] → Deliver [key deliverables] → Address competitor gaps in [specific areas].

        ### Phase 2: Weeks 5-12 (Strategic Enhancement)
        Advanced skills in [advanced areas] → Complete [strategic projects] → Achieve parity with top competitors.

        ## Competitive Intelligence
        
        {"**Benchmark Analysis**: Based on analysis of " + str(skill_gap_data.get('competitor_count', 0)) + " competitor companies" if skill_gap_data else "**Internal Assessment**: Based on role requirements and industry standards"}
        {"**Key Competitors**: " + ', '.join(skill_gap_data.get('competitor_companies', [])) if skill_gap_data else ""}
        {"**Assessment Confidence**: " + str(skill_gap_data.get('evidence_flags', {}).get('confidence_score', 'N/A') * 100) + "%" if skill_gap_data and skill_gap_data.get('evidence_flags', {}).get('confidence_score') else ""}
        
        ### Priority Recommendations
        {"".join(['- ' + rec + chr(10) for rec in skill_gap_data.get('recommendations', [])]) if skill_gap_data else "- Focus on selected priority skills for maximum impact"}

        Instructions:
        1. Analyze the employee's missing skills from prioritySkillsSelected and skill gap data
        2. Map course chapters to specific skill gaps with priority levels
        3. Create realistic POCs based on their role, existing skills, and competitor benchmarks
        4. Focus on business impact relevant to their company/industry and competitive positioning
        5. Ensure timeline is practical and progressive
        6. Use technical tools and frameworks appropriate for their field
        7. Make POCs specific and measurable with competitive context
        8. Connect learning to direct business value and competitive advantage
        9. Include competitor benchmark information where available
        10. Prioritize critical gaps over nice-to-have skills
        """

        try:
            response = self.client.chat.completions.create(
                model=self.deployment_id,
                messages=[
                    {"role": "system", "content": "You are an expert learning and development specialist who creates detailed skill development plans with competitive intelligence. Always follow the exact format requested and provide practical, actionable recommendations that consider competitive benchmarking data."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=3000,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"Error generating plan for {employee_data['employeeName']}: {e}")
            return None
    
    def markdown_to_word(self, markdown_content: str, employee_name: str) -> Document:
        """Convert markdown content to Word document with proper formatting"""
        doc = Document()
        
        lines = markdown_content.split('\n')
        table_rows = []
        in_table = False
        current_headers = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Handle headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], 1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            
            # Handle table
            elif line.startswith('|') and ('Timeline' in line or 'Skill Gap' in line):
                # Table header found
                in_table = True
                current_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                continue
            elif line.startswith('|') and '---' in line:
                # Table separator, skip
                continue
            elif line.startswith('|') and in_table:
                # Table row
                row_data = [cell.strip() for cell in line.split('|')[1:-1]]
                table_rows.append(row_data)
            elif in_table and not line.startswith('|'):
                # End of table, create it
                if table_rows and current_headers:
                    table = doc.add_table(rows=1, cols=len(current_headers))
                    table.style = 'Table Grid'
                    
                    # Add headers
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(current_headers):
                        if i < len(header_cells):
                            header_cells[i].text = header
                            # Make header bold
                            for paragraph in header_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                    
                    # Add data rows
                    for row_data in table_rows:
                        row_cells = table.add_row().cells
                        for i, cell_data in enumerate(row_data):
                            if i < len(row_cells):
                                row_cells[i].text = cell_data
                
                table_rows = []
                in_table = False
                current_headers = []
                
                # Continue with current line (probably paragraph text)
                if line and not line.startswith('#'):
                    doc.add_paragraph(line)
            
            # Handle regular paragraphs and lists
            elif not line.startswith('#') and not line.startswith('|') and not in_table:
                if line.startswith('- **') or line.startswith('**'):
                    # Bold text paragraph
                    p = doc.add_paragraph()
                    if line.startswith('- '):
                        line = line[2:]  # Remove bullet
                    # Simple bold formatting
                    run = p.add_run(line)
                    if '**' in line:
                        # Basic bold parsing
                        parts = line.split('**')
                        p.clear()
                        for i, part in enumerate(parts):
                            run = p.add_run(part)
                            if i % 2 == 1:  # Odd indices should be bold
                                run.bold = True
                elif line.startswith('- '):
                    # Regular bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    # Numbered list item
                    doc.add_paragraph(line, style='List Number')
                else:
                    doc.add_paragraph(line)
        
        # Handle any remaining table data
        if table_rows and in_table and current_headers:
            table = doc.add_table(rows=1, cols=len(current_headers))
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(current_headers):
                if i < len(header_cells):
                    header_cells[i].text = header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            
            # Add data rows
            for row_data in table_rows:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = cell_data
        
        return doc
    
    def save_word_report(self, employee_name: str, report_content: str, output_dir: str = "reports"):
        """Save the report as a Word document"""
        # Create output directory if it doesn't exist
        Path(output_dir).mkdir(exist_ok=True)
        
        # Clean employee name for filename
        clean_name = "".join(c for c in employee_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_name = clean_name.replace(' ', '_')
        
        filename = f"{clean_name}_Enhanced_0-3_Month_Skill_Boost_Plan.docx"
        filepath = Path(output_dir) / filename
        
        try:
            # Convert markdown to Word document
            doc = self.markdown_to_word(report_content, employee_name)
            
            # Save the document
            doc.save(filepath)
            print(f"✓ Enhanced Word report saved: {filepath}")
            return str(filepath)
        except Exception as e:
            print(f"Error saving Word report for {employee_name}: {e}")
            return None
    
    def process_all_employees(self, json_file_path: str, skill_gaps_file: str = "final_skill_gaps_detailed_gpt.json", output_dir: str = "reports"):
        """Process all employees and generate individual enhanced Word reports"""
        
        # Load the course data
        course_data = self.load_course_data(json_file_path)
        
        # Load the skill gaps data
        skill_gaps_data = self.load_skill_gaps_data(skill_gaps_file)
        
        if not course_data:
            print("No course data loaded. Please check your file path and format.")
            return
        
        total_employees = len(course_data)
        print(f"Found {total_employees} employees to process")
        print(f"Skill gaps data loaded for {len(skill_gaps_data)} employees")
        
        successful_reports = []
        failed_reports = []
        
        for i, employee in enumerate(course_data, 1):
            employee_name = employee.get('employeeName', f'Employee_{i}')
            print(f"\n[{i}/{total_employees}] Processing: {employee_name}")
            
            # Find corresponding skill gap data
            skill_gap_info = self.find_skill_gap_data(employee_name, skill_gaps_data)
            if skill_gap_info:
                print(f"  ✓ Found skill gap data - {skill_gap_info.get('competitor_count', 0)} competitors analyzed")
            else:
                print(f"  ⚠ No skill gap data found for {employee_name}")
            
            # Generate the enhanced report
            report_content = self.generate_skill_boost_plan(employee, skill_gap_info)
            
            if report_content:
                # Save the report as Word document
                saved_path = self.save_word_report(employee_name, report_content, output_dir)
                if saved_path:
                    successful_reports.append((employee_name, saved_path))
                else:
                    failed_reports.append(employee_name)
            else:
                failed_reports.append(employee_name)
            
            # Add delay to avoid rate limiting
            if i < total_employees:
                time.sleep(2)
        
        # Summary
        print(f"\n" + "="*60)
        print(f"PROCESSING COMPLETE")
        print(f"="*60)
        print(f"Total Employees: {total_employees}")
        print(f"Enhanced Reports Generated: {len(successful_reports)}")
        print(f"Failed Reports: {len(failed_reports)}")
        
        if successful_reports:
            print(f"\nSuccessfully generated enhanced Word reports:")
            for name, path in successful_reports:
                print(f"  • {name}: {path}")
        
        if failed_reports:
            print(f"\nFailed to generate reports for:")
            for name in failed_reports:
                print(f"  • {name}")

def main():
    """Main function to run the script"""
    
    print("Enhanced Course to Skill Boost Plan Generator (Azure OpenAI + Word Export)")
    print("="*75)
    print("This version includes competitive benchmarking and skill gap analysis")
    print()
   
    # Get input file path
    input_file = input("Enter path to course JSON file (or press Enter for 'spectre_courses.json'): ").strip()
    if not input_file:
        input_file = "spectre_courses.json"
    
    # Get skill gaps file path
    skill_gaps_file = input("Enter path to skill gaps JSON file (or press Enter for 'final_skill_gaps_detailed_gpt.json'): ").strip()
    if not skill_gaps_file:
        skill_gaps_file = "final_skill_gaps_detailed_gpt.json"
    
    # Get output directory
    output_dir = input("Enter output directory (or press Enter for 'reports'): ").strip()
    if not output_dir:
        output_dir = "reports"
    
    # Check if files exist
    if not os.path.exists(input_file):
        print(f"❌ Course file not found: {input_file}")
        return
    
    if not os.path.exists(skill_gaps_file):
        print(f"⚠ Skill gaps file not found: {skill_gaps_file}")
        print("Will proceed with course data only...")
    
    try:
        # Initialize generator
        generator = SkillBoostPlanGenerator()
        
        # Process all employees with enhanced data
        generator.process_all_employees(input_file, skill_gaps_file, output_dir)
        
    except Exception as e:
        print(f"Error initializing generator: {e}")

if __name__ == "__main__":
    main()